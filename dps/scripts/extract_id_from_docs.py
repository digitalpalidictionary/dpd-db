#!/usr/bin/env python3

"""
This script reads IDs from tables in a Word document, removes duplicates, and writes the unique IDs into a CSV file.
"""

from docx import Document
import csv
import sys
from rich.console import Console
from tools.tic_toc import tic, toc

from dps.tools.paths_dps import DPSPaths

console = Console()


def extract_ids_from_docx(filename):
    """
    Extracts IDs from tables inside a Word document.
    
    :param filename: Path to the .docx file.
    :return: List of extracted IDs.
    """
    doc = Document(filename)
    all_ids = []

    for table in doc.tables:
        for row in table.rows:
            # Assuming the ID is in the first cell of every row
            cell_content = row.cells[0].text.strip()
            # Check if the cell content is an integer
            try:
                _id = int(cell_content)
                all_ids.append(_id)
            except ValueError:
                # The content is not a number, so we skip it
                continue

    return all_ids

def write_ids_to_csv(ids, output_file):
    """
    Writes a list of IDs to a CSV file.
    
    :param ids: List of IDs.
    :param output_file: Path to the output CSV file.
    """
    with open(output_file, 'w') as f:
        writer = csv.writer(f, delimiter='\t')
        writer.writerow(['id'])  # Header row
        for _id in ids:
            writer.writerow([_id])


def convert_docx_to_txt(dpspth, input_docx_file):
    doc = Document(input_docx_file)
    
    # Extract all the text from the document
    doc_text = []
    for paragraph in doc.paragraphs:
        doc_text.append(paragraph.text)

    output_txt_file = dpspth.text_to_add_path
    with open(output_txt_file, 'w') as f:
        f.write('\n'.join(doc_text))

    console.print(f"[bold green]Text extracted from {input_docx_file} written to {output_txt_file}")


def main():

    tic()

    console.print("[bold blue]Enter the name of the .docx file (without extension):")
    file_name = input("").strip()
    dpspth = DPSPaths()
    if dpspth.local_downloads_dir:
        input_docx_file = f"{dpspth.local_downloads_dir}/{file_name}.docx"
    else:
        console.print("[bold red]no path to Downloads")
        sys.exit()
    if input_docx_file:
        convert_docx_to_txt(dpspth, input_docx_file)
    output_csv_file = dpspth.id_to_add_path
    
    ids = extract_ids_from_docx(input_docx_file)
    unique_ids = list(set(ids))  # Convert to a set to get unique IDs, then convert back to a list
    write_ids_to_csv(ids, output_csv_file)
    console.print(f"[bold green]IDs written to {output_csv_file}")
    console.print(f"[bold yellow]from {file_name}.docx")
    console.print(f"Number of rows of unique IDs extracted: [bold]{len(unique_ids)}[/bold]")


    toc()

if __name__ == "__main__":
    main()
